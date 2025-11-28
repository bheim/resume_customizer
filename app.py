import os, hashlib, base64, tempfile
from typing import Optional, List
from fastapi import FastAPI, UploadFile, File, Form, Body, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel
from config import log, health, supabase
from docx_utils import load_docx, collect_word_numbered_bullets, set_paragraph_text_with_selective_links, enforce_single_page
from llm_utils import rewrite_with_openai, generate_followup_questions, should_ask_more_questions, rewrite_with_context
from caps import tiered_char_cap, enforce_char_cap_with_reprompt
from scoring import composite_score
from db_utils import (create_qa_session, get_qa_session, store_qa_pair, update_qa_answer,
                      get_session_qa_pairs, get_user_context, store_user_context,
                      update_session_status, get_answered_qa_pairs)


# Helper function for robust hex/base64 decoding
def decode_base64(data: str) -> bytes:
    """
    Decode hex or base64-encoded string with comprehensive error handling.

    Primary format: Hex-encoded (consistent with Supabase BYTEA storage/retrieval)
    Fallback format: Base64-encoded (for backward compatibility)

    This function provides bulletproof decoding for data retrieved from
    Supabase BYTEA columns. All resume data is now stored as hex-encoded
    for consistency with Supabase's native BYTEA return format.

    Handles:
    - Hex-encoded format (PRIMARY): "\\x504b03..." - PostgreSQL hex format
    - Base64-encoded format (FALLBACK): "UEsDBB..." - for legacy data
    - Invalid padding (not multiple of 4) - automatically adds '=' padding
    - Embedded whitespace (newlines, tabs, spaces) - strips all whitespace
    - Leading/trailing whitespace - removed
    - Invalid characters - validates and reports which characters are invalid

    Args:
        data (str): Encoded string (hex or base64, may contain whitespace)

    Returns:
        bytes: Decoded binary data

    Raises:
        ValueError: If data contains invalid characters for both formats
        binascii.Error: If decoding fails after all fixes

    History:
        - v1: Basic base64 decoding with strip()
        - v2: Added automatic padding (fixes "not multiple of 4" error)
        - v3: Added comprehensive whitespace removal and character validation
        - v4: Added hex-encoding support for Supabase BYTEA columns
        - v5: Switched to hex-encoding as primary format throughout application
    """
    import re

    # Remove ALL whitespace (spaces, tabs, newlines, etc.)
    data = re.sub(r'\s+', '', data)

    # Detect format: hex-encoded (PRIMARY) or base64 (FALLBACK)
    if data.startswith('\\x') or '\\x' in data[:20]:
        # Hex-encoded format - PRIMARY/EXPECTED format for all resume storage
        log.debug(f"Decoding hex-encoded data (length: {len(data)} chars)")
        try:
            # Parse hex string with \xHH format
            # Replace \x with nothing and decode pairs of hex digits
            hex_str = data.replace('\\x', '')

            # Validate it's valid hex
            if not re.match(r'^[0-9a-fA-F]*$', hex_str):
                invalid_chars = set(re.findall(r'[^0-9a-fA-F]', hex_str))
                log.error(f"Invalid hex characters: {invalid_chars}")
                raise ValueError(f"Invalid hex characters: {invalid_chars}")

            # Decode hex to bytes
            result = bytes.fromhex(hex_str)
            log.debug(f"Successfully decoded {len(hex_str)} hex chars to {len(result)} bytes")
            return result

        except Exception as e:
            log.error(f"Hex decode failed")
            log.error(f"Data length: {len(data)}, First 100 chars: {data[:100]}")
            log.error(f"Error type: {type(e).__name__}, Message: {str(e)}")
            raise

    else:
        # Base64-encoded format - FALLBACK for legacy data or edge cases
        log.info(f"Decoding base64-encoded data (length: {len(data)} chars) - legacy format")

        # Validate that string only contains valid base64 characters
        # Valid: A-Z, a-z, 0-9, +, /, = (padding)
        if not re.match(r'^[A-Za-z0-9+/]*=*$', data):
            # Find invalid characters for debugging
            invalid_chars = set(re.findall(r'[^A-Za-z0-9+/=]', data))
            log.error(f"Base64 decode failed - Invalid characters found: {invalid_chars}")
            log.error(f"Data length: {len(data)}, First 100 chars: {data[:100]}")
            raise ValueError(f"Invalid base64 characters: {invalid_chars}")

        # Add padding if needed (must be multiple of 4)
        missing_padding = len(data) % 4
        if missing_padding:
            data += '=' * (4 - missing_padding)
            log.debug(f"Added {4 - missing_padding} padding characters to base64 string")

        try:
            result = base64.b64decode(data)
            log.debug(f"Successfully decoded {len(data)} base64 chars to {len(result)} bytes")
            return result
        except Exception as e:
            # Detailed error logging for debugging
            log.error(f"Base64 decode failed after validation and padding")
            log.error(f"Data length: {len(data)} (multiple of 4: {len(data) % 4 == 0})")
            log.error(f"First 100 chars: {data[:100]}")
            log.error(f"Last 20 chars: {data[-20:] if len(data) > 20 else data}")
            log.error(f"Error type: {type(e).__name__}, Message: {str(e)}")
            raise


# Import new v2 endpoints
try:
    from api_endpoints_new import router as v2_router
    v2_endpoints_available = True
except Exception as e:
    log.warning(f"Could not import v2 endpoints: {e}")
    v2_endpoints_available = False

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["POST", "GET", "OPTIONS", "DELETE"],
    allow_headers=["*"],
    expose_headers=["X-Score-Before", "X-Score-After", "X-Score-Delta", "X-QA-Context-Used", "X-Session-Id"]
)

# Include v2 router if available
# DISABLED: Using simpler v2 endpoints defined below instead
# if v2_endpoints_available:
#     app.include_router(v2_router)
#     log.info("✓ V2 endpoints registered")


# Pydantic models for Q&A endpoints
class AnswerSubmission(BaseModel):
    session_id: str
    answers: List[dict]  # List of {"qa_id": "...", "answer": "..."}
    user_id: Optional[str] = None


# Pydantic models for v2 context endpoints
class ContextStartRequest(BaseModel):
    user_id: str
    bullet_text: str


class ContextAnswerRequest(BaseModel):
    session_id: str
    user_answer: str


class ConfirmFactsRequest(BaseModel):
    session_id: str
    facts: dict
    user_confirmed: bool = True


# Pydantic models for v2 apply endpoints
class BulletItem(BaseModel):
    bullet_text: str
    bullet_id: Optional[str] = None
    use_stored_facts: bool = True


class BulletGenerationRequest(BaseModel):
    user_id: str
    job_description: str
    bullets: List[BulletItem]


class EnhancedBulletItem(BaseModel):
    original: str
    enhanced: str
    used_facts: bool


class BulletGenerationResponse(BaseModel):
    enhanced_bullets: List[EnhancedBulletItem]


# Pydantic models for base resume endpoints
class BaseResumeUploadResponse(BaseModel):
    success: bool
    file_name: str
    message: str


class BaseResumeInfoResponse(BaseModel):
    has_resume: bool
    file_name: Optional[str] = None
    uploaded_at: Optional[str] = None


@app.get("/")
def root():
    return health()


@app.post("/v2/context/start")
async def v2_context_start(request: ContextStartRequest):
    """
    Simpler conversational context endpoint matching Lovable's expectations.
    Returns a single initial question.
    """
    from llm_utils import generate_conversational_question

    log.info(f"/v2/context/start called: user_id={request.user_id}, bullet_text={request.bullet_text[:100]}")

    try:
        # Create session
        session_id = create_qa_session(request.user_id, "", [request.bullet_text])
        if not session_id:
            return JSONResponse({"error": "failed_to_create_session"}, status_code=500)

        # Generate first question
        initial_question = generate_conversational_question(request.bullet_text)

        log.info(f"Generated initial question for session {session_id}: {initial_question[:100]}")

        return JSONResponse({
            "session_id": session_id,
            "initial_question": initial_question,
            "message": "Session started. Answer the question or skip."
        })

    except Exception as e:
        log.exception(f"Error in /v2/context/start: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/v2/context/answer")
async def v2_context_answer(request: ContextAnswerRequest):
    """
    Simpler answer submission endpoint matching Lovable's expectations.
    Takes a single answer, returns either next question or extracted facts.
    """
    log.info(f"/v2/context/answer called: session_id={request.session_id}, answer_length={len(request.user_answer)}")
    log.info(f"Answer content: {request.user_answer[:200]}")

    try:
        # Get session
        session = get_qa_session(request.session_id)
        if not session:
            return JSONResponse({"error": "session_not_found"}, status_code=404)

        bullets = session.get("bullets", [])
        bullet_text = bullets[0] if bullets else ""

        # Store this answer (create a new Q&A pair for it)
        qa_id = store_qa_pair(request.session_id, "Question", request.user_answer, "conversational", 0)

        # Get all answers so far
        answered_qa = get_answered_qa_pairs(request.session_id)
        log.info(f"Total answers collected: {len(answered_qa)}")

        # Check if we have enough context
        from llm_utils import should_ask_more_questions, extract_facts_from_conversation

        # Build conversation history
        conversation_history = "\n\n".join([
            f"Q: {qa['question']}\nA: {qa['answer']}"
            for qa in answered_qa
        ])

        need_more, reason = should_ask_more_questions(answered_qa, bullets, "")
        log.info(f"Need more questions: {need_more}, reason: {reason}")

        if not need_more or len(answered_qa) >= 5:  # Max 5 questions
            # Extract facts from conversation
            log.info("Extracting facts from conversation")
            facts = extract_facts_from_conversation(bullet_text, conversation_history)

            return JSONResponse({
                "status": "complete",
                "extracted_facts": facts,
                "message": "Context gathering complete!"
            })
        else:
            # Generate next question
            from llm_utils import generate_conversational_question
            next_question = generate_conversational_question(bullet_text)

            return JSONResponse({
                "status": "continue",
                "next_question": next_question,
                "conversation_so_far": conversation_history
            })

    except Exception as e:
        log.exception(f"Error in /v2/context/answer: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/v2/context/confirm_facts")
async def v2_confirm_facts(request: ConfirmFactsRequest):
    """
    Save confirmed facts for a bullet.
    """
    log.info(f"/v2/context/confirm_facts called: session_id={request.session_id}")

    try:
        # Get session to find bullet info
        session = get_qa_session(request.session_id)
        if not session:
            return JSONResponse({"error": "session_not_found"}, status_code=404)

        user_id = session.get("user_id")
        bullets = session.get("bullets", [])
        bullet_text = bullets[0] if bullets else ""

        # Store bullet with facts
        from llm_utils import embed
        from db_utils import store_user_bullet, store_bullet_facts

        embedding = embed(bullet_text)
        bullet_id = store_user_bullet(user_id, bullet_text, embedding)

        if bullet_id:
            fact_id = store_bullet_facts(bullet_id, request.facts, request.session_id, request.user_confirmed)
            log.info(f"Stored facts for bullet {bullet_id}, fact_id={fact_id}")

            return JSONResponse({
                "bullet_id": bullet_id,
                "message": "Facts saved successfully"
            })
        else:
            return JSONResponse({"error": "failed_to_store_bullet"}, status_code=500)

    except Exception as e:
        log.exception(f"Error in /v2/context/confirm_facts: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/v2/bullets/{user_id}")
async def get_user_bullets(user_id: str):
    """
    Get all bullets for a user with their associated facts.
    Returns bullets with facts for the "My Bullets" page.
    """
    log.info(f"/v2/bullets/{user_id} called")

    try:
        from db_utils import get_bullet_facts

        # Query user_bullets table
        if not supabase:
            return JSONResponse({"error": "supabase_not_configured"}, status_code=503)

        # Get all bullets for this user
        result = supabase.table('user_bullets').select(
            'id, bullet_text, created_at, updated_at, source_resume_name'
        ).eq('user_id', user_id).order('created_at', desc=True).execute()

        bullets_data = []
        for bullet in result.data:
            bullet_id = bullet['id']

            # Get facts for this bullet
            facts_list = get_bullet_facts(bullet_id, confirmed_only=False)

            # Get the most recent facts
            latest_facts = None
            has_confirmed_facts = False
            if facts_list:
                latest_facts = facts_list[0].get('facts')
                has_confirmed_facts = facts_list[0].get('confirmed_by_user', False)

            bullets_data.append({
                "bullet_id": bullet_id,
                "bullet_text": bullet['bullet_text'],
                "has_facts": latest_facts is not None,
                "has_confirmed_facts": has_confirmed_facts,
                "facts": latest_facts,
                "created_at": bullet['created_at'],
                "updated_at": bullet['updated_at'],
                "source_resume": bullet.get('source_resume_name')
            })

        log.info(f"Found {len(bullets_data)} bullets for user {user_id}")

        return JSONResponse({
            "bullets": bullets_data,
            "count": len(bullets_data)
        })

    except Exception as e:
        log.exception(f"Error getting user bullets: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# =====================================================================
# BASE RESUME ENDPOINTS
# =====================================================================

@app.post("/v2/user/base_resume")
async def upload_base_resume(
    user_id: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Upload and store a base resume for the user.
    This resume can be reused across multiple job applications.
    """
    log.info(f"/v2/user/base_resume POST called for user {user_id}")

    if not supabase:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    try:
        # Read file
        file_content = await file.read()
        file_name = file.filename or "base_resume.docx"

        # Validate it's a valid DOCX
        try:
            doc = load_docx(file_content)
            bullets, _ = collect_word_numbered_bullets(doc)
            log.info(f"Validated DOCX with {len(bullets)} bullets")
        except Exception as e:
            log.exception("Invalid DOCX file")
            raise HTTPException(status_code=400, detail=f"Invalid DOCX file: {str(e)}")

        # Hex-encode bytes for BYTEA column storage (consistent with Supabase return format)
        log.info(f"Upload - File size: {len(file_content)} bytes")
        file_data_hex = '\\x' + file_content.hex()

        # Store in database (upsert)
        data = {
            "user_id": user_id,
            "file_name": file_name,
            "file_data": file_data_hex,  # Store hex-encoded string
            "updated_at": "now()"
        }

        result = supabase.table("user_base_resumes").upsert(
            data,
            on_conflict="user_id"
        ).execute()

        log.info(f"Stored base resume for user {user_id}: {file_name} ({len(file_content)} bytes)")

        # Immediately verify what was stored
        verify = supabase.table("user_base_resumes").select("file_data").eq("user_id", user_id).execute()
        if verify.data:
            stored_data_b64 = verify.data[0]["file_data"]
            stored_data = decode_base64(stored_data_b64)
            log.info(f"Verify - Stored data length: {len(stored_data)} bytes")
            log.info(f"Verify - Match original: {stored_data == file_content}")
            if stored_data != file_content:
                log.error("CORRUPTION DETECTED: Stored data doesn't match original!")
                log.error(f"Original length: {len(file_content)}, Stored length: {len(stored_data)}")
                log.error(f"Difference: {len(stored_data) - len(file_content)} bytes")

        return {
            "success": True,
            "file_name": file_name,
            "message": "Base resume uploaded successfully"
        }

    except HTTPException:
        raise
    except Exception as e:
        log.exception(f"Error uploading base resume: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/v2/user/base_resume/{user_id}")
async def get_base_resume_info(user_id: str):
    """
    Get information about the user's base resume (if it exists).
    Returns metadata only, not the actual file.
    """
    log.info(f"/v2/user/base_resume GET called for user {user_id}")

    if not supabase:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    try:
        result = supabase.table("user_base_resumes").select(
            "file_name, uploaded_at, updated_at"
        ).eq("user_id", user_id).execute()

        if result.data and len(result.data) > 0:
            resume = result.data[0]
            return {
                "has_resume": True,
                "file_name": resume["file_name"],
                "uploaded_at": resume.get("updated_at") or resume.get("uploaded_at")
            }
        else:
            return {
                "has_resume": False
            }

    except Exception as e:
        log.exception(f"Error getting base resume info: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/v2/user/base_resume/{user_id}")
async def delete_base_resume(user_id: str):
    """
    Delete the user's base resume.
    """
    log.info(f"/v2/user/base_resume DELETE called for user {user_id}")

    if not supabase:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    try:
        result = supabase.table("user_base_resumes").delete().eq(
            "user_id", user_id
        ).execute()

        log.info(f"Deleted base resume for user {user_id}")

        return {
            "success": True,
            "message": "Base resume deleted successfully"
        }

    except Exception as e:
        log.exception(f"Error deleting base resume: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =====================================================================
# JOB APPLICATION ENDPOINTS
# =====================================================================

@app.post("/v2/apply/match_bullets")
async def match_bullets_for_job(
    user_id: str = Form(...),
    resume_file: Optional[UploadFile] = File(None)
):
    """
    Match bullets from uploaded resume to stored bullets with facts.
    This is the first step in the job application flow.

    If resume_file is not provided, uses the user's stored base resume.
    Returns a session_id that can be used to retrieve this resume later.
    """
    log.info(f"/v2/apply/match_bullets called for user {user_id}")

    try:
        import uuid
        from db_utils_optimized import match_bullet_with_confidence_optimized
        from llm_utils import embed
        from db_utils import get_bullet_facts
        from docx import Document

        # Generate session ID for this job application
        session_id = str(uuid.uuid4())

        # Get resume content (either from upload or base resume)
        if resume_file:
            log.info("Using uploaded resume file")
            content = await resume_file.read()
            resume_name = resume_file.filename or "resume.docx"

            # Store uploaded resume in session table for later retrieval
            if not supabase:
                raise HTTPException(status_code=503, detail="Supabase not configured")

            try:
                # Encode as hex for storage (consistent with base resume format)
                content_hex = '\\x' + content.hex()

                supabase.table("job_application_sessions").insert({
                    "user_id": user_id,
                    "session_id": session_id,
                    "resume_data": content_hex,
                    "resume_name": resume_name
                }).execute()
                log.info(f"Stored session resume for session {session_id}")
            except Exception as e:
                log.warning(f"Failed to store session resume: {e}")
                # Continue anyway - not critical
        else:
            log.info("No file uploaded, fetching base resume")
            if not supabase:
                raise HTTPException(status_code=503, detail="Supabase not configured")

            result = supabase.table("user_base_resumes").select("file_data").eq("user_id", user_id).execute()

            if not result.data or len(result.data) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="No resume file provided and no base resume found. Please upload a resume or set a base resume."
                )

            # Decode base64 string from BYTEA column to get raw bytes
            content_b64 = result.data[0]["file_data"]
            content = decode_base64(content_b64)

            # Log the retrieval for debugging
            log.info(f"Retrieval - Data length: {len(content)} bytes")
            log.info(f"Retrieval - Data type: {type(content)}")
            log.info(f"Loaded base resume from database ({len(content)} bytes)")

        # Extract bullets from resume
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        doc = Document(tmp_path)
        bullets, _ = collect_word_numbered_bullets(doc)
        os.unlink(tmp_path)

        if not bullets:
            raise HTTPException(status_code=400, detail="No bullets found in resume")

        log.info(f"Extracted {len(bullets)} bullets from resume")

        # Match each bullet
        matches = []
        for idx, bullet in enumerate(bullets):
            embedding = embed(bullet)
            match_result = match_bullet_with_confidence_optimized(user_id, bullet, embedding)

            # Get facts if match found
            facts = None
            if match_result["bullet_id"]:
                fact_records = get_bullet_facts(match_result["bullet_id"], confirmed_only=True)
                if fact_records:
                    facts = fact_records[0]["facts"]

            matches.append({
                "bullet_index": idx,
                "bullet_text": bullet,
                **match_result,
                "has_facts": facts is not None,
                "facts": facts
            })

            log.info(f"Bullet {idx}: matched={match_result['bullet_id'] is not None}, confidence={match_result.get('confidence', 0):.2f}, has_facts={facts is not None}")

        log.info(f"Matched {sum(1 for m in matches if m['bullet_id'])} out of {len(bullets)} bullets")

        return JSONResponse({
            "session_id": session_id,
            "bullets": bullets,
            "matches": matches
        })

    except HTTPException:
        raise
    except Exception as e:
        log.exception(f"Error matching bullets: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v2/apply/generate_with_facts")
async def generate_resume_with_facts(request: BulletGenerationRequest):
    """
    Generate optimized resume using stored facts for matched bullets.
    This is the second step in the job application flow.
    """
    log.info(f"/v2/apply/generate_with_facts called for user {request.user_id} with {len(request.bullets)} bullets")

    try:
        from db_utils_optimized import match_bullet_with_confidence_optimized
        from llm_utils import embed, generate_bullet_with_facts
        from db_utils import get_bullet_facts

        # For each bullet, try to find stored facts
        enhanced_bullets = []
        with_facts_count = 0
        without_facts_count = 0

        for idx, bullet_item in enumerate(request.bullets):
            bullet_text = bullet_item.bullet_text
            used_facts = False
            enhanced_text = bullet_text  # Default to original

            # Skip fact-based generation if explicitly requested
            if not bullet_item.use_stored_facts:
                log.info(f"Bullet {idx} opted out of using stored facts")
            else:
                # Use provided bullet_id if available, otherwise match
                bullet_id = bullet_item.bullet_id
                if not bullet_id:
                    embedding = embed(bullet_text)
                    match_result = match_bullet_with_confidence_optimized(
                        request.user_id,
                        bullet_text,
                        embedding
                    )
                    bullet_id = match_result.get("bullet_id")

                # Get facts if we have a bullet_id
                facts = None
                if bullet_id:
                    fact_records = get_bullet_facts(bullet_id, confirmed_only=True)
                    if fact_records:
                        facts = fact_records[0]["facts"]

                if facts:
                    # Generate with facts
                    log.info(f"Generating bullet {idx} with stored facts (bullet_id: {bullet_id})")
                    enhanced_text = generate_bullet_with_facts(
                        bullet_text,
                        request.job_description,
                        facts
                    )
                    used_facts = True
                    with_facts_count += 1
                else:
                    # Generate without facts (still optimize using job description)
                    log.info(f"Bullet {idx} has no stored facts, optimizing without facts")
                    enhanced_text = generate_bullet_with_facts(
                        bullet_text,
                        request.job_description,
                        {}  # Empty facts dict
                    )
                    used_facts = False
                    without_facts_count += 1

            # Add to results
            enhanced_bullets.append({
                "original": bullet_text,
                "enhanced": enhanced_text,
                "used_facts": used_facts
            })

        log.info(f"Generated {with_facts_count} bullets with facts, {without_facts_count} without facts")

        return {
            "enhanced_bullets": enhanced_bullets
        }

    except Exception as e:
        log.exception(f"Error generating with facts: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/download")
async def download_resume(
    bullets: Optional[str] = Form(None),
    user_id: Optional[str] = Form(None),
    session_id: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    """
    Generate final resume DOCX with enhanced bullets.

    Priority order for resume source:
    1. Uploaded file (if provided)
    2. Session resume (if session_id provided)
    3. Base resume (if exists)

    Returns modified DOCX file for download.
    """
    log.info(f"/download called")
    log.info(f"Received parameters: bullets={'present' if bullets else 'missing'}, user_id={user_id}, session_id={session_id}, file={'present' if file else 'missing'}")

    # Validate required parameters
    if not bullets:
        log.error("Missing required parameter: bullets")
        raise HTTPException(status_code=400, detail="Missing required parameter: bullets (JSON string)")

    if not user_id:
        log.error("Missing required parameter: user_id")
        raise HTTPException(status_code=400, detail="Missing required parameter: user_id")

    try:
        import json

        # Parse bullets from JSON string
        try:
            enhanced_bullets = json.loads(bullets)
            if not isinstance(enhanced_bullets, list):
                raise HTTPException(status_code=400, detail="bullets must be a JSON array")
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"Invalid JSON in bullets: {str(e)}")

        log.info(f"Received {len(enhanced_bullets)} enhanced bullets")

        # Extract enhanced text from bullets (handle both string and object formats)
        enhanced_texts = []
        for bullet in enhanced_bullets:
            if isinstance(bullet, dict):
                # Frontend sent {original, enhanced, used_facts} format
                enhanced_texts.append(bullet.get("enhanced", bullet.get("original", "")))
            else:
                # Frontend sent just strings
                enhanced_texts.append(str(bullet))

        log.info(f"Extracted {len(enhanced_texts)} enhanced bullet texts")

        # Get resume content with priority: uploaded file > session resume > base resume
        if file:
            log.info("Priority 1: Using uploaded resume file")
            raw = await file.read()
            file_name = file.filename or "resume.docx"
        elif session_id:
            log.info(f"Priority 2: Fetching session resume for session {session_id}")
            if not supabase:
                raise HTTPException(status_code=503, detail="Supabase not configured")

            result = supabase.table("job_application_sessions").select(
                "resume_data, resume_name"
            ).eq("session_id", session_id).execute()

            if result.data and len(result.data) > 0:
                # Decode from base64
                raw = decode_base64(result.data[0]["resume_data"])
                file_name = result.data[0]["resume_name"]
                log.info(f"Loaded session resume: {file_name}")
            else:
                # Session not found, fall back to base resume
                log.warning(f"Session {session_id} not found, falling back to base resume")
                result = supabase.table("user_base_resumes").select(
                    "file_data, file_name"
                ).eq("user_id", user_id).execute()

                if not result.data or len(result.data) == 0:
                    raise HTTPException(
                        status_code=400,
                        detail="Session resume not found and no base resume exists. Please upload a resume."
                    )

                # Get raw bytes from BYTEA column
                raw = result.data[0]["file_data"]
                file_name = result.data[0]["file_name"]
                log.info(f"Loaded base resume: {file_name}")
        else:
            log.info("Priority 3: Fetching base resume")
            if not supabase:
                raise HTTPException(status_code=503, detail="Supabase not configured")

            result = supabase.table("user_base_resumes").select(
                "file_data, file_name"
            ).eq("user_id", user_id).execute()

            if not result.data or len(result.data) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="No resume file, session, or base resume found. Please upload a resume."
                )

            # Decode from base64
            raw = decode_base64(result.data[0]["file_data"])
            file_name = result.data[0]["file_name"]
            log.info(f"Loaded base resume: {file_name}")

        # Validate file
        size = len(raw)
        log.info(f"Using file='{file_name}' size={size}")

        if not raw or size < 512:
            raise HTTPException(status_code=400, detail="File is empty or too small")

        # Parse DOCX (this validates it's a valid DOCX file)
        try:
            doc = load_docx(raw)
        except Exception as e:
            log.exception("Failed to parse DOCX")
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {str(e)}")

        bullets_in_doc, paras = collect_word_numbered_bullets(doc)
        if not bullets_in_doc:
            raise HTTPException(status_code=422, detail="No bullets found in resume")

        log.info(f"Found {len(bullets_in_doc)} bullets in document")

        if len(enhanced_texts) != len(paras):
            raise HTTPException(
                status_code=400,
                detail=f"Bullet count mismatch: document has {len(paras)}, received {len(enhanced_texts)}"
            )

        # Update document with enhanced bullets
        for idx, (p, orig, new_text) in enumerate(zip(paras, bullets_in_doc, enhanced_texts)):
            # Apply character cap
            cap = tiered_char_cap(len(orig))
            fitted = enforce_char_cap_with_reprompt(new_text, cap)
            set_paragraph_text_with_selective_links(p, fitted)

        # Enforce single page layout
        enforce_single_page(doc)

        # Save modified document to bytes
        from io import BytesIO
        buf = BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        log.info(f"Generated DOCX file with {len(enhanced_texts)} enhanced bullets")

        # Return the file as a download
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": 'attachment; filename="resume_optimized.docx"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        log.exception(f"Error in /download: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload")
async def upload(file: UploadFile = File(...), user_id: str = Form(...)):
    """
    Upload resume and extract bullets.
    Simple endpoint that returns just the bullets from the resume.
    """
    # Read and validate file
    raw = await file.read()
    size = len(raw)
    ct = file.content_type
    log.info(f"/upload recv file='{file.filename}' size={size} user_id={user_id}")

    if not raw or size < 512:
        return JSONResponse({"error": "empty_or_small_file"}, status_code=400)

    if ct not in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                  "application/octet-stream", "application/msword"}:
        return JSONResponse({"error": "bad_content_type", "got": ct}, status_code=415)

    # Parse DOCX
    try:
        doc = load_docx(raw)
    except Exception as e:
        log.exception("bad_docx")
        return JSONResponse({"error": "bad_docx", "detail": str(e)}, status_code=400)

    bullets, paras = collect_word_numbered_bullets(doc)

    # Log all paragraphs for debugging
    log.info(f"Document has {len(doc.paragraphs)} paragraphs total")
    for i, p in enumerate(doc.paragraphs[:10]):  # Log first 10 paragraphs
        log.info(f"Para {i}: '{p.text[:100]}'")

    if not bullets:
        log.warning("No bullets found using Word numbering or bullet chars: •·-–—◦●*")
        return JSONResponse({"error": "no_bullets_found", "detail": "Resume must use numbered lists or bullet points (•·-–—◦●*)"}, status_code=422)

    log.info(f"Extracted {len(bullets)} bullets for user {user_id}")
    log.info(f"Sample bullets: {[b[:60] for b in bullets[:3]]}")

    return JSONResponse({
        "bullets": bullets,
        "message": f"Successfully extracted {len(bullets)} bullets"
    })


@app.post("/rewrite")
async def rewrite(file: UploadFile = File(...), job_description: str = Form(...), max_chars_override: Optional[int] = Form(None)):
    raw = await file.read()
    size = len(raw); ct = file.content_type; sha = hashlib.sha256(raw).hexdigest()
    if not raw or size < 512: return JSONResponse({"error":"empty_or_small_file"}, status_code=400)
    if ct not in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document","application/octet-stream","application/msword"}:
        return JSONResponse({"error":"bad_content_type","got":ct}, status_code=415)
    try:
        doc = load_docx(raw)
    except Exception as e:
        log.exception("bad_docx"); return JSONResponse({"error":"bad_docx","detail":str(e)}, status_code=400)

    bullets, paras = collect_word_numbered_bullets(doc)
    if not bullets: return JSONResponse({"error":"no_bullets_found"}, status_code=422)
    log.info(f"bullets_in={len(bullets)} sample_in={[b[:60] for b in bullets[:3]]}")

    try:
        rewritten = rewrite_with_openai(bullets, job_description)
    except Exception as e:
        log.exception("openai_failed"); return JSONResponse({"error":"openai_failed","detail":str(e)}, status_code=502)

    log.info(f"bullets_out={len(rewritten)} sample_out={[r[:60] for r in rewritten[:3]]}")
    if len(rewritten) != len(paras):
        return JSONResponse({"error":"bullet_count_mismatch","in":len(paras),"out":len(rewritten)}, status_code=500)

    final_texts = []
    for idx, (p, orig, new_text) in enumerate(zip(paras, bullets, rewritten)):
        cap = tiered_char_cap(len(orig), max_chars_override)
        fitted = enforce_char_cap_with_reprompt(new_text, cap)
        set_paragraph_text_with_selective_links(p, fitted)
        final_texts.append(fitted)

    enforce_single_page(doc)
    from io import BytesIO
    buf = BytesIO(); doc.save(buf); data = buf.getvalue()
    return Response(content=data, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": 'attachment; filename="resume_edited.docx"'})

@app.post("/rewrite_json")
async def rewrite_json(file: UploadFile = File(...), job_description: str = Form(...), max_chars_override: Optional[int] = Form(None)):
    raw = await file.read()
    size = len(raw); ct = file.content_type; sha = hashlib.sha256(raw).hexdigest()
    log.info(f"/rewrite_json recv file='{file.filename}' size={size} sha256={sha}")
    if not raw or size < 512: return JSONResponse({"error":"empty_or_small_file"}, status_code=400)
    if ct not in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document","application/octet-stream","application/msword"}:
        return JSONResponse({"error":"bad_content_type","got":ct}, status_code=415)
    try:
        doc = load_docx(raw)
    except Exception as e:
        log.exception("bad_docx"); return JSONResponse({"error":"bad_docx","detail":str(e)}, status_code=400)

    bullets, paras = collect_word_numbered_bullets(doc)
    if not bullets: return JSONResponse({"error":"no_bullets_found"}, status_code=422)

    resume_before = "\n".join(bullets)
    try:
        score_before = composite_score(resume_before, job_description)
    except Exception as e:
        score_before = {"embed_sim":0.0,"keyword_cov":0.0,"llm_score":0.0,"composite":0.0,"error":str(e)}

    try:
        rewritten = rewrite_with_openai(bullets, job_description)
    except Exception as e:
        log.exception("openai_failed"); return JSONResponse({"error":"openai_failed","detail":str(e)}, status_code=502)

    if len(rewritten) != len(paras):
        return JSONResponse({"error":"bullet_count_mismatch","in":len(paras),"out":len(rewritten)}, status_code=500)

    final_texts = []
    for idx, (p, orig, new_text) in enumerate(zip(paras, bullets, rewritten)):
        cap = tiered_char_cap(len(orig), max_chars_override)
        fitted = enforce_char_cap_with_reprompt(new_text, cap)
        set_paragraph_text_with_selective_links(p, fitted)
        final_texts.append(fitted)

    resume_after = "\n".join(final_texts)
    try:
        score_after = composite_score(resume_after, job_description)
    except Exception as e:
        score_after = {"embed_sim":0.0,"keyword_cov":0.0,"llm_score":0.0,"composite":0.0,"error":str(e)}

    enforce_single_page(doc)
    from io import BytesIO
    buf = BytesIO(); doc.save(buf); data = buf.getvalue()
    b64 = base64.b64encode(data).decode("ascii")

    delta = {}
    try:
        delta = {
            "embed_sim": round(score_after["embed_sim"] - score_before["embed_sim"], 4),
            "keyword_cov": round(score_after["keyword_cov"] - score_before["keyword_cov"], 4),
            "llm_score": round(score_after["llm_score"] - score_before["llm_score"], 1),
            "composite": round(score_after["composite"] - score_before["composite"], 1),
        }
    except Exception: pass

    return JSONResponse({
        "file_b64": b64,
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "filename": "resume_edited.docx",
        "scores": {"before": score_before, "after": score_after, "delta": delta}
    })


@app.post("/generate_questions")
async def generate_questions(file: UploadFile = File(...), job_description: str = Form(...), user_id: Optional[str] = Form(None)):
    """
    Generate initial follow-up questions based on resume bullets and job description.
    Creates a Q&A session in Supabase and returns questions for the frontend to display.
    """
    if not supabase:
        return JSONResponse({"error": "supabase_not_configured", "detail": "Supabase is not configured. Set SUPABASE_URL and SUPABASE_KEY."}, status_code=503)

    # Read and validate file
    raw = await file.read()
    size = len(raw)
    ct = file.content_type
    log.info(f"/generate_questions recv file='{file.filename}' size={size}")

    if not raw or size < 512:
        return JSONResponse({"error": "empty_or_small_file"}, status_code=400)

    if ct not in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                  "application/octet-stream", "application/msword"}:
        return JSONResponse({"error": "bad_content_type", "got": ct}, status_code=415)

    # Parse DOCX
    try:
        doc = load_docx(raw)
    except Exception as e:
        log.exception("bad_docx")
        return JSONResponse({"error": "bad_docx", "detail": str(e)}, status_code=400)

    bullets, paras = collect_word_numbered_bullets(doc)
    if not bullets:
        return JSONResponse({"error": "no_bullets_found"}, status_code=422)

    log.info(f"Found {len(bullets)} bullets")

    # Check for existing user context to avoid repeat questions
    existing_context = []
    if user_id and user_id != "anonymous":
        user_context = get_user_context(user_id)
        existing_context = [{"question": ctx["question"], "answer": ctx["answer"]} for ctx in user_context]
        log.info(f"Found {len(existing_context)} existing Q&A pairs for user {user_id}")
        if existing_context:
            log.info(f"Existing questions for context: {[qa['question'] for qa in existing_context[:5]]}")  # Log first 5
    else:
        log.info(f"No user_id provided or anonymous user, skipping context retrieval")

    # Create Q&A session
    session_id = create_qa_session(user_id, job_description, bullets)
    if not session_id:
        return JSONResponse({"error": "failed_to_create_session"}, status_code=500)

    # Generate questions using LLM (let LLM decide how many, up to 10)
    try:
        questions = generate_followup_questions(bullets, job_description, existing_context, max_questions=10)
    except Exception as e:
        log.exception("Failed to generate questions")
        return JSONResponse({"error": "question_generation_failed", "detail": str(e)}, status_code=502)

    # Store questions in database
    qa_pairs = []
    for q in questions:
        qa_id = store_qa_pair(session_id, q["question"], question_type=q["type"], bullet_index=q.get("bullet_index"))
        if qa_id:
            log.info(f"Created qa_pair with ID: {qa_id} for question: {q['question'][:50]}...")
            qa_pairs.append({
                "qa_id": qa_id,
                "question": q["question"],
                "type": q["type"],
                "bullet_index": q.get("bullet_index", 0),
                "bullet_text": q.get("bullet_text", "")
            })
        else:
            log.error(f"Failed to store qa_pair for question: {q['question'][:50]}...")

    log.info(f"Generated {len(qa_pairs)} questions for session {session_id}")
    log.info(f"Returning qa_ids to frontend: {[qp['qa_id'] for qp in qa_pairs]}")

    return JSONResponse({
        "session_id": session_id,
        "questions": qa_pairs,
        "bullet_count": len(bullets)
    })


@app.post("/submit_answers")
async def submit_answers(submission: AnswerSubmission = Body(...)):
    """
    Receive answers to follow-up questions and determine if more questions are needed.
    Stores answers in Supabase and optionally stores in user context.
    """
    if not supabase:
        return JSONResponse({"error": "supabase_not_configured"}, status_code=503)

    session_id = submission.session_id
    answers = submission.answers
    user_id = submission.user_id

    log.info(f"Received answers for session {session_id}")
    log.info(f"qa_ids received from frontend: {[ans.get('qa_id') for ans in answers]}")

    # Validate session exists
    session = get_qa_session(session_id)
    if not session:
        return JSONResponse({"error": "session_not_found"}, status_code=404)

    # Update answers in database
    for ans in answers:
        qa_id = ans.get("qa_id")
        answer_text = ans.get("answer", "").strip()

        if not qa_id or not answer_text:
            continue

        success = update_qa_answer(qa_id, answer_text)
        if not success:
            log.warning(f"Failed to update answer for qa_id={qa_id}")

    # Get all answered Q&A pairs
    all_qa_pairs = get_session_qa_pairs(session_id)
    answered_qa = [{"question": qa["question"], "answer": qa["answer"]}
                   for qa in all_qa_pairs if qa.get("answer")]

    log.info(f"Session {session_id} now has {len(answered_qa)} answered questions")

    # Store in user context if user_id provided
    if user_id and user_id != "anonymous":
        log.info(f"Storing {len(answered_qa)} Q&A pairs in user_context for user {user_id}")
        for qa in answered_qa:
            success = store_user_context(user_id, qa["question"], qa["answer"])
            if success:
                log.info(f"Stored in user_context: {qa['question'][:50]}...")
            else:
                log.warning(f"Failed to store in user_context: {qa['question'][:50]}...")
    else:
        log.info(f"No user_id or anonymous - not storing in user_context")

    # Determine if more questions are needed
    bullets = session["bullets"]
    job_description = session["job_description"]

    try:
        need_more, reason = should_ask_more_questions(answered_qa, bullets, job_description)
    except Exception as e:
        log.exception("Failed to determine if more questions needed")
        need_more = False
        reason = "Error determining question need"

    log.info(f"Need more questions: {need_more}. Reason: {reason}")

    # If more questions needed, generate them
    new_questions = []
    if need_more:
        try:
            questions = generate_followup_questions(bullets, job_description, answered_qa, max_questions=10)

            for q in questions:
                qa_id = store_qa_pair(session_id, q["question"], question_type=q["type"], bullet_index=q.get("bullet_index"))
                if qa_id:
                    new_questions.append({
                        "qa_id": qa_id,
                        "question": q["question"],
                        "type": q["type"],
                        "bullet_index": q.get("bullet_index", 0),
                        "bullet_text": q.get("bullet_text", "")
                    })
        except Exception as e:
            log.exception("Failed to generate follow-up questions")

    # If no more questions needed, mark session as ready for rewriting
    if not need_more:
        update_session_status(session_id, "ready_for_rewrite")

    return JSONResponse({
        "session_id": session_id,
        "need_more_questions": need_more,
        "reason": reason,
        "new_questions": new_questions,
        "total_answered": len(answered_qa),
        "ready_for_rewrite": not need_more
    })


@app.post("/rewrite_with_qa")
async def rewrite_with_qa(session_id: str = Form(...), max_chars_override: Optional[int] = Form(None)):
    """
    Rewrite resume using both the job description and the Q&A context from the session.
    This endpoint should be called after the Q&A flow is complete.
    """
    if not supabase:
        return JSONResponse({"error": "supabase_not_configured"}, status_code=503)

    # Validate session exists
    session = get_qa_session(session_id)
    if not session:
        return JSONResponse({"error": "session_not_found"}, status_code=404)

    bullets = session["bullets"]
    job_description = session["job_description"]

    # Get answered Q&A pairs for context
    qa_context = get_answered_qa_pairs(session_id)

    if not qa_context:
        log.warning(f"No Q&A context found for session {session_id}, falling back to basic rewrite")
        # Fallback to basic rewrite without context
        try:
            rewritten = rewrite_with_openai(bullets, job_description)
        except Exception as e:
            log.exception("openai_failed")
            return JSONResponse({"error": "openai_failed", "detail": str(e)}, status_code=502)
    else:
        # Use Q&A context for better rewriting
        log.info(f"Using {len(qa_context)} Q&A pairs for context in rewrite")
        try:
            rewritten = rewrite_with_context(bullets, job_description, qa_context)
        except Exception as e:
            log.exception("openai_failed_with_context")
            return JSONResponse({"error": "openai_failed", "detail": str(e)}, status_code=502)

    # We need the original document to modify it
    # For now, return just the rewritten bullets as JSON
    # The frontend will need to provide the original file or we need to store it

    resume_before = "\n".join(bullets)
    try:
        score_before = composite_score(resume_before, job_description)
    except Exception as e:
        score_before = {"embed_sim": 0.0, "keyword_cov": 0.0, "llm_score": 0.0, "composite": 0.0, "error": str(e)}

    resume_after = "\n".join(rewritten)
    try:
        score_after = composite_score(resume_after, job_description)
    except Exception as e:
        score_after = {"embed_sim": 0.0, "keyword_cov": 0.0, "llm_score": 0.0, "composite": 0.0, "error": str(e)}

    delta = {}
    try:
        delta = {
            "embed_sim": round(score_after["embed_sim"] - score_before["embed_sim"], 4),
            "keyword_cov": round(score_after["keyword_cov"] - score_before["keyword_cov"], 4),
            "llm_score": round(score_after["llm_score"] - score_before["llm_score"], 1),
            "composite": round(score_after["composite"] - score_before["composite"], 1),
        }
    except Exception:
        pass

    # Mark session as completed
    update_session_status(session_id, "completed")

    return JSONResponse({
        "session_id": session_id,
        "original_bullets": bullets,
        "rewritten_bullets": rewritten,
        "scores": {"before": score_before, "after": score_after, "delta": delta},
        "qa_context_used": len(qa_context)
    })


@app.post("/generate_results")
async def generate_results(file: UploadFile = File(...), session_id: str = Form(...), max_chars_override: Optional[int] = Form(None)):
    """
    Generate final resume DOCX with bullets rewritten using Q&A context.
    Accepts the original resume file and session_id, returns modified DOCX for download.
    """
    log.info(f"Received generate_results request for session: {session_id}")

    if not supabase:
        return JSONResponse({"error": "supabase_not_configured"}, status_code=503)

    # Read and validate file
    raw = await file.read()
    size = len(raw)
    ct = file.content_type
    log.info(f"generate_results recv file='{file.filename}' size={size}")

    if not raw or size < 512:
        return JSONResponse({"error": "empty_or_small_file"}, status_code=400)

    if ct not in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                  "application/octet-stream", "application/msword"}:
        return JSONResponse({"error": "bad_content_type", "got": ct}, status_code=415)

    # Parse DOCX
    try:
        doc = load_docx(raw)
    except Exception as e:
        log.exception("bad_docx")
        return JSONResponse({"error": "bad_docx", "detail": str(e)}, status_code=400)

    bullets, paras = collect_word_numbered_bullets(doc)
    if not bullets:
        return JSONResponse({"error": "no_bullets_found"}, status_code=422)

    log.info(f"Found {len(bullets)} bullets in uploaded file")

    # Validate session exists
    session = get_qa_session(session_id)
    if not session:
        return JSONResponse({"error": "session_not_found"}, status_code=404)

    job_description = session["job_description"]

    # Calculate before score
    resume_before = "\n".join(bullets)
    try:
        score_before = composite_score(resume_before, job_description)
        log.info(f"Score BEFORE: {score_before}")
    except Exception as e:
        log.exception(f"Failed to calculate score_before: {e}")
        score_before = {"embed_sim": 0.0, "keyword_cov": 0.0, "llm_score": 0.0, "composite": 0.0, "error": str(e)}

    # Get answered Q&A pairs for context
    qa_context = get_answered_qa_pairs(session_id)

    if not qa_context:
        log.warning(f"No Q&A context found for session {session_id}, falling back to basic rewrite")
        try:
            rewritten = rewrite_with_openai(bullets, job_description)
        except Exception as e:
            log.exception("openai_failed")
            return JSONResponse({"error": "openai_failed", "detail": str(e)}, status_code=502)
    else:
        log.info(f"Using {len(qa_context)} Q&A pairs for context in rewrite")
        try:
            rewritten = rewrite_with_context(bullets, job_description, qa_context)
        except Exception as e:
            log.exception("openai_failed_with_context")
            return JSONResponse({"error": "openai_failed", "detail": str(e)}, status_code=502)

    log.info(f"Rewritten {len(rewritten)} bullets")

    if len(rewritten) != len(paras):
        return JSONResponse({"error": "bullet_count_mismatch", "in": len(paras), "out": len(rewritten)}, status_code=500)

    # Update document with rewritten bullets, enforcing character limits
    final_texts = []
    for idx, (p, orig, new_text) in enumerate(zip(paras, bullets, rewritten)):
        cap = tiered_char_cap(len(orig), max_chars_override)
        fitted = enforce_char_cap_with_reprompt(new_text, cap)
        set_paragraph_text_with_selective_links(p, fitted)
        final_texts.append(fitted)

    # Calculate after score
    resume_after = "\n".join(final_texts)
    try:
        score_after = composite_score(resume_after, job_description)
        log.info(f"Score AFTER: {score_after}")
    except Exception as e:
        log.exception(f"Failed to calculate score_after: {e}")
        score_after = {"embed_sim": 0.0, "keyword_cov": 0.0, "llm_score": 0.0, "composite": 0.0, "error": str(e)}

    # Calculate deltas
    delta = {}
    try:
        delta = {
            "embed_sim": round(score_after["embed_sim"] - score_before["embed_sim"], 4),
            "keyword_cov": round(score_after["keyword_cov"] - score_before["keyword_cov"], 4),
            "llm_score": round(score_after["llm_score"] - score_before["llm_score"], 1),
            "composite": round(score_after["composite"] - score_before["composite"], 1),
        }
        log.info(f"Score DELTA: {delta}")
    except Exception as e:
        log.exception(f"Failed to calculate delta: {e}")
        pass

    # Enforce single page layout
    enforce_single_page(doc)

    # Save modified document to bytes
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    # Mark session as completed
    update_session_status(session_id, "completed")

    log.info(f"Returning modified DOCX for session {session_id}")

    # Return the file as a download with scores in custom headers
    # This avoids JSON parsing issues with large base64 strings
    import json

    # Log what we're sending in headers
    headers_to_send = {
        "Content-Disposition": 'attachment; filename="resume_customized.docx"',
        "X-Session-Id": session_id,
        "X-Score-Before": json.dumps(score_before),
        "X-Score-After": json.dumps(score_after),
        "X-Score-Delta": json.dumps(delta),
        "X-QA-Context-Used": str(len(qa_context))
    }
    log.info(f"Response headers being sent:")
    log.info(f"  X-Score-Before: {headers_to_send['X-Score-Before']}")
    log.info(f"  X-Score-After: {headers_to_send['X-Score-After']}")
    log.info(f"  X-Score-Delta: {headers_to_send['X-Score-Delta']}")
    log.info(f"  X-QA-Context-Used: {headers_to_send['X-QA-Context-Used']}")

    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers_to_send
    )
