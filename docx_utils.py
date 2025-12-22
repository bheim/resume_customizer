from copy import deepcopy
from io import BytesIO
from typing import List, Tuple
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.section import WD_SECTION_START
from config import BULLET_CHARS

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def load_docx(raw: bytes) -> Document:
    return Document(BytesIO(raw))

def _collect_links(p):
    links, p_el = [], p._p
    for h in p_el.findall(f".//{{{W_NS}}}hyperlink"):
        r_texts, rPr_template = [], None
        for r in h.findall(f".//{{{W_NS}}}r"):
            t = r.find(f".//{{{W_NS}}}t")
            if t is not None and t.text:
                r_texts.append(t.text)
            rPr = r.find(f"./{{{W_NS}}}rPr")
            if rPr is not None and rPr_template is None:
                rPr_template = deepcopy(rPr)
        anchor_text = "".join(r_texts).strip()
        r_id = h.get(qn("r:id")); url = None
        if r_id:
            try:
                rel = p.part.rels[r_id]
                if rel.is_external and rel.reltype == RT.HYPERLINK:
                    url = rel.target_ref
            except KeyError:
                pass
        if anchor_text and url:
            links.append({"text": anchor_text, "url": url, "rPr": rPr_template})
    return links

def _make_run(text, rPr_template=None):
    r = OxmlElement("w:r")
    if rPr_template is not None:
        r.append(deepcopy(rPr_template))
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t); return r

def _make_hyperlink_run(p, text, url, rPr_template=None):
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    h.append(_make_run(text, rPr_template)); return h

def set_paragraph_text_with_selective_links(p, new_text):
    p_el = p._p
    tmpl = None
    for r in p_el.findall(f"./{{{W_NS}}}r"):
        tmpl = r.find(f"./{{{W_NS}}}rPr")
        if tmpl is not None:
            tmpl = deepcopy(tmpl)
        break
    links = _collect_links(p)
    anchors = sorted(links, key=lambda d: len(d["text"]), reverse=True)
    lower = new_text.lower(); spans=[]; i=0
    while i < len(new_text):
        match=None
        for lk in anchors:
            a=lk["text"]; al=a.lower()
            if lower.startswith(al, i):
                match=(i, i+len(a), lk); break
        if match:
            spans.append(match); i=match[1]
        else:
            next_pos=len(new_text)
            for lk in anchors:
                j=lower.find(lk["text"].lower(), i)
                if j!=-1: next_pos=min(next_pos, j)
            spans.append((i, next_pos, None)); i=next_pos
    for child in list(p_el):
        if child.tag != f"{{{W_NS}}}pPr":
            p_el.remove(child)
    for start,end,lk in spans:
        seg=new_text[start:end]
        if not seg: continue
        p_el.append(_make_hyperlink_run(p, seg, lk["url"], lk["rPr"] or tmpl) if lk else _make_run(seg, tmpl))

def enforce_single_page(doc: Document):
    body = doc.element.body
    for p in doc.paragraphs:
        for br in list(p._element.findall(f".//{{{W_NS}}}br")):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)
        pPr = p._p.pPr
        if pPr is not None:
            for pb in list(pPr.findall(f"./{{{W_NS}}}pageBreakBefore")):
                pPr.remove(pb)
    for sectPr in body.findall(f".//{{{W_NS}}}sectPr"):
        t = sectPr.find(f"./{{{W_NS}}}type") or OxmlElement("w:type")
        t.set(qn("w:val"), "continuous")
        if t.getparent() is None: sectPr.append(t)
    try:
        if doc.sections:
            doc.sections[-1].start_type = WD_SECTION_START.CONTINUOUS
    except Exception:
        pass
    def _body_p(): return body.findall(f"./{{{W_NS}}}p")
    while len(_body_p())>1 and not doc.paragraphs[-1].text.strip():
        body.remove(doc.paragraphs[-1]._element)

def collect_word_numbered_bullets(doc: Document, use_heuristics: bool = False) -> Tuple[List[str], List]:
    """
    Collect bullet points from a Word document.

    Args:
        doc: Word Document object
        use_heuristics: If True, use lenient heuristics for PDF-converted docs

    Returns:
        Tuple of (bullet_texts, paragraph_objects)
    """
    bullets, paras = [], []
    def _is_numbered(p):
        pPr = p._p.pPr
        return (
            (pPr is not None)
            and (getattr(pPr, "numPr", None) is not None)
            and (getattr(pPr.numPr, "numId", None) is not None)
            and (getattr(pPr.numPr.numId, "val", None) is not None)
        )

    def _looks_like_bullet(text: str) -> bool:
        """Heuristic check if text looks like a bullet point."""
        if not text or len(text) < 10:  # Too short
            return False

        # Check if starts with bullet char
        if text[0] in BULLET_CHARS:
            return True

        # Check if starts with common patterns (for PDF-converted docs)
        # e.g., "• ", "- ", "o ", "▪ ", etc.
        import re
        bullet_patterns = [
            r'^[•·\-–—◦●*○▪▫■□➢➣⚫⚪]\s',  # Bullet chars with space
            r'^\d+[\.\)]\s',  # Numbers like "1. " or "1) "
        ]
        for pattern in bullet_patterns:
            if re.match(pattern, text):
                return True

        return False

    # Process paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t: continue

        is_glyph = t and t[0] in BULLET_CHARS
        is_numbered_list = _is_numbered(p)

        # For PDF-converted docs, also check heuristics
        if use_heuristics and not is_numbered_list and not is_glyph:
            if _looks_like_bullet(t):
                # Extract text after bullet char
                import re
                match = re.match(r'^[•·\-–—◦●*○▪▫■□➢➣⚫⚪]\s+(.+)$', t)
                if match:
                    t = match.group(1)
                    bullets.append(t)
                    paras.append(p)
                    continue
                match = re.match(r'^\d+[\.\)]\s+(.+)$', t)
                if match:
                    t = match.group(1)
                    bullets.append(t)
                    paras.append(p)
                    continue

        # Standard bullet detection
        if is_numbered_list or is_glyph:
            if is_glyph: t = t[1:].lstrip()
            bullets.append(t)
            paras.append(p)

    # Process tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if not t: continue
                    is_glyph = t and t[0] in BULLET_CHARS
                    if _is_numbered(p) or is_glyph:
                        if is_glyph: t = t[1:].lstrip()
                        bullets.append(t); paras.append(p)
    return bullets, paras