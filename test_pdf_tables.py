#!/usr/bin/env python3
"""Check if bullet text is in tables."""

from docx import Document

doc = Document("pdfres_converted.docx")

print("=" * 80)
print(f"TABLES IN DOCUMENT: {len(doc.tables)}")
print("=" * 80)
print()

for table_idx, table in enumerate(doc.tables):
    print(f"\nTABLE {table_idx + 1}:")
    print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")

    for row_idx, row in enumerate(table.rows[:10]):  # First 10 rows
        print(f"\n  Row {row_idx + 1}:")
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"    Cell {cell_idx + 1}: '{text[:100]}'")
                # Check paragraphs in cell
                for p_idx, p in enumerate(cell.paragraphs):
                    p_text = p.text.strip()
                    if p_text:
                        print(f"      Para {p_idx + 1}: '{p_text[:80]}'")
